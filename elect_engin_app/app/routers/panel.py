from __future__ import annotations
from typing import Optional, Dict, Union
from fastapi import APIRouter, HTTPException, BackgroundTasks, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, ValidationError
from pathlib import Path
import tempfile, shutil, zipfile, os
import logging

logger = logging.getLogger(__name__)

# --- Internal project imports ---
from app.schemas.panel_ir import PanelScheduleIR                  # IR schema for panel schedule
from app.io.panel_excel import write_excel_from_ir                # IR → Excel (.xlsx), preserves template formatting
from app.export.pdf import export_pdf_from_ir                     # IR → PDF
from app.routers.preflight import _kva_formulas_per_phase         # Server-side fallback KVA formulas

# --- FastAPI Router Setup ---
router = APIRouter(prefix="/panel", tags=["panel"])

# --- Key paths ---
ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "out"
OUT.mkdir(parents=True, exist_ok=True)
TEMPLATE_XLSX = ROOT / "templates" / "panelboard_template.xlsx"


# -------------------------------------------------------------------
# MODELS
# -------------------------------------------------------------------
class ExportPayload(BaseModel):
    """
    Preferred request body for /panel/export/zip:
      {
        "ir": { ...PanelScheduleIR... },
        "_kva_formulas": { "G57": "=G56*277/1000", ... },   # optional (from GPT)
        "_inferred_system": "3PH WYE, GROUNDED"             # optional (from GPT)
      }

    Aliases let us use underscore keys in JSON while accessing them as
    .kva_formulas / .inferred_system in Python.
    """
    ir: PanelScheduleIR
    kva_formulas: Optional[Dict[str, str]] = Field(default=None, alias="_kva_formulas")
    inferred_system: Optional[str] = Field(default=None, alias="_inferred_system")

    model_config = {
        "populate_by_name": True,      # allow using field names or aliases
        "protected_namespaces": (),    # permit alias keys that start with '_' in JSON
    }


# -------------------------------------------------------------------
# ENDPOINTS
# -------------------------------------------------------------------

@router.post("/export/zip")
def export_zip(
    payload: Union[ExportPayload, PanelScheduleIR],
    background_tasks: BackgroundTasks,
    confirm: bool = Query(False),
    skip_ai_review: bool = Query(False),
    session: str = Query(None),
):
    """
    Builds an Excel + PDF panel schedule and returns a ZIP with optional AI technical review.
    - Accepts EITHER:
        1) ExportPayload (preferred): { ir, _kva_formulas?, _inferred_system? }
        2) Raw PanelScheduleIR (legacy clients)
    - Uses GPT-provided formulas when present; otherwise falls back to server logic.
    - `skip_ai_review=true` to skip the OpenAI technical review (default: false)
    - AI review is ADVISORY ONLY - never blocks builds. Results included in ZIP.
    """
    # ---- 1) Extract & validate IR; capture GPT extras if present ----
    if isinstance(payload, PanelScheduleIR):
        ir = payload
        gpt_formulas = None
        # inferred_system unused server-side for now, but kept for future logging/telemetry
    else:
        ir = payload.ir
        gpt_formulas = payload.kva_formulas
        # payload.inferred_system is available if you want to log/emit it

    try:
        ir = PanelScheduleIR.model_validate(ir)
    except ValidationError as e:
        raise HTTPException(status_code=400, detail=e.errors())

    # ---- 2) Guard: ensure template exists ----
    if not TEMPLATE_XLSX.exists():
        raise HTTPException(
            status_code=500,
            detail=f"Master template not found at {TEMPLATE_XLSX}. "
                   "Place 'panelboard_template.xlsx' under /templates."
        )

    # ---- 3) Choose formulas (prefer GPT; otherwise compute) ----
    formulas = gpt_formulas or _kva_formulas_per_phase(ir)

    # (Optional) enforce confirmation on server side in the future:
    # if not confirm:
    #     raise HTTPException(status_code=409, detail="Build not confirmed by user.")

    # ---- 4) Get task-specific output directory ----
    from app.main import get_task_directories
    _, outputs_dir = get_task_directories(session)
    if not outputs_dir:
        raise HTTPException(404, "No active task found. Please start a task before building.")
    
    # ---- 5) Build in a temporary workdir ----
    workdir = Path(tempfile.mkdtemp())
    try:
        placeholder_xlsx = workdir / "panel_schedule.xlsx"

        # Excel writer returns the final saved path inside outputs_dir with your filename convention
        excel_real_path = write_excel_from_ir(
            ir=ir,
            out_path=str(placeholder_xlsx),
            template_xlsx=str(TEMPLATE_XLSX),
            formulas=formulas,
            outputs_dir=outputs_dir,
        )

        # Matching PDF next to the Excel
        pdf_path = excel_real_path.with_suffix(".pdf")
        export_pdf_from_ir(ir=ir, out_pdf=str(pdf_path))

        # ---- AI Technical Review (Advisory Only - Never Blocks) ----
        review_data = None
        review_path = None
        txt_path = None
        docx_path = None
        if not skip_ai_review:
            try:
                from app.ai.gpt_preflight import run_gpt_preflight
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import json
                
                review_data = run_gpt_preflight(ir)
                
                # Save JSON and TXT to backend directory (not visible to user in outputs)
                from app.main import get_task_directories
                uploads_dir, _ = get_task_directories(session)
                backend_dir = uploads_dir.parent / "backend" if uploads_dir else outputs_dir.parent / "backend"
                backend_dir.mkdir(exist_ok=True)
                
                # Save review as JSON file (backend only)
                review_filename = excel_real_path.with_suffix("").name + "_AI_REVIEW.json"
                review_path = backend_dir / review_filename
                with open(review_path, 'w') as f:
                    json.dump(review_data, f, indent=2)
                
                # Create Word document with AI review (visible to user)
                docx_filename = excel_real_path.with_suffix("").name + "_AI_REVIEW.docx"
                docx_path = outputs_dir / docx_filename
                
                doc = Document()
                
                # Title
                title = doc.add_heading('AI ELECTRICAL ENGINEERING REVIEW', level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Subtitle
                subtitle = doc.add_paragraph('(ADVISORY ONLY)')
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_run = subtitle.runs[0]
                subtitle_run.font.size = Pt(10)
                subtitle_run.font.italic = True
                
                doc.add_paragraph()  # Spacing
                
                # Panel info
                panel_info = doc.add_paragraph()
                panel_info.add_run(f'Panel: ').bold = True
                panel_info.add_run(ir.header.panel_name)
                
                status_para = doc.add_paragraph()
                status_para.add_run('Review Status: ').bold = True
                status_text = '⚠️ ISSUES FOUND' if not review_data.get('ok_to_build') else '✓ APPROVED'
                status_run = status_para.add_run(status_text)
                if not review_data.get('ok_to_build'):
                    status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                else:
                    status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                
                doc.add_paragraph()  # Spacing
                
                # Summary
                if review_data.get('summary'):
                    doc.add_heading('SUMMARY', level=2)
                    doc.add_paragraph(review_data['summary'])
                
                # System Analysis
                if review_data.get('system_analysis'):
                    doc.add_heading('SYSTEM ANALYSIS', level=2)
                    for key, value in review_data['system_analysis'].items():
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(value))
                
                # Warnings
                if review_data.get('warnings'):
                    doc.add_heading('⚠️ WARNINGS', level=2)
                    for w in review_data['warnings']:
                        p = doc.add_paragraph(w, style='List Bullet')
                        p.runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange
                
                # Recommendations
                if review_data.get('recommendations'):
                    doc.add_heading('💡 RECOMMENDATIONS', level=2)
                    for r in review_data['recommendations']:
                        doc.add_paragraph(r, style='List Bullet')
                
                # Detailed Checklist
                doc.add_heading('DETAILED REVIEW CHECKLIST', level=2)
                for item in review_data.get('items', []):
                    status = "✓" if item.get('pass') else "✗"
                    p = doc.add_paragraph()
                    status_run = p.add_run(f"{status} ")
                    status_run.bold = True
                    if item.get('pass'):
                        status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    else:
                        status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    
                    p.add_run(item.get('check', 'N/A'))
                    
                    if item.get('notes'):
                        notes_p = doc.add_paragraph(f"   Notes: {item['notes']}", style='List Bullet 2')
                        notes_p.runs[0].font.italic = True
                
                # Footer note
                doc.add_paragraph()  # Spacing
                footer = doc.add_paragraph()
                footer.add_run('NOTE: ').bold = True
                footer.add_run('This review is advisory only. Build files are generated regardless of findings. ')
                footer.add_run('Review the findings, update your design as needed, then rebuild.')
                footer.runs[0].font.size = Pt(9)
                footer.runs[1].font.size = Pt(9)
                footer.runs[2].font.size = Pt(9)
                
                doc.save(str(docx_path))
                
                # Also create TXT summary (backend only - not visible to user)
                txt_filename = excel_real_path.with_suffix("").name + "_AI_REVIEW.txt"
                txt_path = backend_dir / txt_filename
                with open(txt_path, 'w') as f:
                    f.write("=" * 80 + "\n")
                    f.write("AI ELECTRICAL ENGINEERING REVIEW (ADVISORY ONLY)\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(f"Panel: {ir.header.panel_name}\n")
                    f.write(f"Review Status: {'⚠️ ISSUES FOUND' if not review_data.get('ok_to_build') else '✓ APPROVED'}\n\n")
                    
                    if review_data.get('summary'):
                        f.write("SUMMARY:\n")
                        f.write(review_data['summary'] + "\n\n")
                    
                    if review_data.get('system_analysis'):
                        f.write("SYSTEM ANALYSIS:\n")
                        for key, value in review_data['system_analysis'].items():
                            f.write(f"  {key.replace('_', ' ').title()}: {value}\n")
                        f.write("\n")
                    
                    if review_data.get('warnings'):
                        f.write("⚠️ WARNINGS:\n")
                        for w in review_data['warnings']:
                            f.write(f"  • {w}\n")
                        f.write("\n")
                    
                    if review_data.get('recommendations'):
                        f.write("💡 RECOMMENDATIONS:\n")
                        for r in review_data['recommendations']:
                            f.write(f"  • {r}\n")
                        f.write("\n")
                    
                    f.write("DETAILED REVIEW:\n")
                    for item in review_data.get('items', []):
                        status = "✓" if item.get('pass') else "✗"
                        f.write(f"{status} {item.get('check', 'N/A')}\n")
                        if item.get('notes'):
                            f.write(f"   Notes: {item['notes']}\n")
                    
                    f.write("\n" + "=" * 80 + "\n")
                    f.write("NOTE: This review is advisory only. Build files are generated regardless.\n")
                    f.write("Review findings and update your design as needed, then rebuild.\n")
                    f.write("=" * 80 + "\n")
            except Exception as e:
                # AI review failed - log but don't block build
                logger.warning(f"AI review failed (non-blocking): {e}", exc_info=True)
                review_data = {
                    "error": str(e),
                    "ok_to_build": True,  # Default to true on error
                    "message": "AI review unavailable - proceeding with build"
                }

        # Zip them together using the Excel base name (user files only)
        zip_name = excel_real_path.with_suffix("").name + ".zip"
        dest_zip = OUT / zip_name
        with zipfile.ZipFile(dest_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(excel_real_path, excel_real_path.name)
            zf.write(pdf_path, pdf_path.name)
            
            # Include Word document (user-facing AI review)
            if docx_path and docx_path.exists():
                zf.write(docx_path, docx_path.name)
    finally:
        shutil.rmtree(workdir, ignore_errors=True)

    # ---- 5) Cleanup the ZIP after response is sent ----
    def _cleanup(path: str):
        try:
            os.remove(path)
        except OSError:
            pass
    background_tasks.add_task(_cleanup, str(dest_zip))

    # ---- 6) Return the ZIP ----
    return FileResponse(
        path=str(dest_zip),
        filename=dest_zip.name,       # e.g., panel_LP-1_208-120V.zip
        media_type="application/zip",
        background=background_tasks,
    )